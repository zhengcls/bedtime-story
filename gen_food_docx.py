# -*- coding: utf-8 -*-
"""Generate a Word document listing foods that can be preserved for 3-5+ years."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Style helpers ---
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def set_run_font(run, name='Microsoft YaHei', size=Pt(10.5), bold=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=Pt(22), bold=True, color=RGBColor(0x1A, 0x5C, 0x2E))

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=Pt(11), color=RGBColor(0x66, 0x66, 0x66))

def add_section_heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=Pt(14), bold=True, color=RGBColor(0x2E, 0x7D, 0x32))

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=Pt(10), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Green background
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2E7D32',
            qn('w:val'): 'clear'
        })
        shading.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=Pt(9.5))
            if c_idx == 0:
                run.bold = True

    # Column widths hint
    widths = [Cm(4), Cm(5.5), Cm(7)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

def add_note(text):
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=Pt(9.5), color=RGBColor(0x55, 0x55, 0x55))

def add_tip_heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=Pt(14), bold=True, color=RGBColor(0xE6, 0x5C, 0x00))

def add_tip(text):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=Pt(10.5))

# ==================== DOCUMENT CONTENT ====================

add_title('可保存3-5年及以上的食物清单')
add_subtitle('按类型详细分类 | 含保存年限与关键条件')
add_subtitle('整理日期：2026年4月')

doc.add_paragraph('')  # spacer

# --- Category 1 ---
add_section_heading('1. 主食 / 谷物类')
add_table(
    ['食物', '实际可保存年限', '关键保存条件'],
    [
        ['大米', '普通密封2-3年；充氮铁罐装可达10年+', '阴凉干燥，充氮密封最佳'],
        ['白面粉', '3-5年', '密封防潮，加入干燥剂'],
        ['燕麦片', '3-5年（未开封真空装）', '避光防潮'],
        ['干意面 / 挂面', '3-5年，过期2年内仍可食用', '干燥密封，防虫'],
        ['藜麦', '3-5年', '避光防潮密封'],
        ['糙米', '1-2年（油脂含量高，不如白米耐存）', '冷藏可延长'],
    ]
)

# --- Category 2 ---
add_section_heading('2. 豆类 / 种子')
add_table(
    ['食物', '实际可保存年限', '关键保存条件'],
    [
        ['干黄豆 / 黑豆 / 红豆 / 绿豆', '3-5年，充氮密封可达10年+', '干燥密封防虫'],
        ['干扁豆 / 鹰嘴豆', '3-5年', '干燥防潮'],
        ['奇亚籽 / 亚麻籽', '3-5年', '避光密封（富含油脂，久存需注意氧化）'],
    ]
)

# --- Category 3 ---
add_section_heading('3. 调味品 / 甜味剂')
add_table(
    ['食物', '实际可保存年限', '关键保存条件'],
    [
        ['蜂蜜', '几乎无限期', '天然低水分抗菌，密封防吸水即可（考古发现3000年前的蜂蜜仍可食用）'],
        ['白砂糖 / 冰糖', '几乎无限期', '密封防潮，不会变质'],
        ['食盐', '几乎无限期', '防潮即可'],
        ['白醋', '5年+', '酸性环境天然抑菌'],
        ['酱油（未开封）', '3-5年', '阴凉避光'],
        ['食用明胶', '几十年', '密封干燥'],
    ]
)

# --- Category 4 ---
add_section_heading('4. 罐头类')
add_table(
    ['食物', '实际可保存年限', '关键保存条件'],
    [
        ['肉类罐头（午餐肉、红烧肉）', '3-5年，军用罐头可达20年+', '密封完好、无鼓包锈蚀即可'],
        ['鱼类罐头（金枪鱼、沙丁鱼）', '3-5年', '高温灭菌密封'],
        ['豆类罐头', '3-5年', '避免高温存放'],
        ['番茄罐头', '3-4年', '酸性较高，略短于肉类罐头'],
        ['水果罐头', '3-5年', '选天然果汁浸泡的'],
    ]
)

# --- Category 5 ---
add_section_heading('5. 干货 / 脱水类')
add_table(
    ['食物', '实际可保存年限', '关键保存条件'],
    [
        ['干香菇 / 干木耳', '3-5年', '密封防潮防虫'],
        ['腐竹 / 粉丝 / 粉条', '3-5年', '干燥密封'],
        ['海带 / 紫菜', '3-5年', '密封防潮'],
        ['脱水蔬菜', '3-5年', '密封干燥'],
        ['冻干水果 / 冻干蔬菜', '5-25年', '真空或密封包装，防潮是关键'],
        ['冻干肉类', '5-25年', '真空或密封包装，防潮是关键'],
        ['红枣 / 葡萄干', '3-5年', '密封阴凉'],
    ]
)

# --- Category 6 ---
add_section_heading('6. 乳制品')
add_table(
    ['食物', '实际可保存年限', '关键保存条件'],
    [
        ['全脂奶粉', '2-3年（密封铁罐可达5年）', '严格防潮，结块即变质'],
        ['脱脂奶粉', '5年+', '脂肪少，更耐存'],
        ['炼乳（未开封罐装）', '3-5年', '阴凉处'],
    ]
)

# --- Category 7 ---
add_section_heading('7. 饮品 / 酒类')
add_table(
    ['食物', '实际可保存年限', '关键保存条件'],
    [
        ['高度白酒（>40度）', '几十年至百年', '密封避光，越陈越香'],
        ['干红葡萄酒', '5年+（好酒更长）', '恒温横放'],
        ['白茶 / 普洱茶', '5-15年+', '密封避光，适度通风'],
        ['速溶咖啡（密封未开封）', '3-5年', '防潮'],
    ]
)

# --- Category 8 ---
add_section_heading('8. 其他')
add_table(
    ['食物', '实际可保存年限', '关键保存条件'],
    [
        ['压缩饼干', '3年+（军粮级20年+）', '多层复合真空密封'],
        ['高温烘烤饼干（铁罐装）', '过期后仍可延长2-10年', '罐装密封'],
        ['黑巧克力（70%+）', '3年+', '密封冷藏'],
        ['瓶装水', '无限期', 'PET瓶密封完好即可'],
    ]
)

# --- Tips section ---
doc.add_paragraph('')
add_tip_heading('四大保存黄金法则')

tips = [
    ('密封', '真空包装/充氮铁罐 > 密封罐 > 原包装，隔绝氧气和湿气是第一要务'),
    ('干燥', '湿度是最大敌人，可在容器内放食品级干燥剂'),
    ('避光', '紫外线加速油脂氧化和营养流失'),
    ('低温', '室温以下越低越好，15\u00b0C以下是理想温度'),
]

for i, (title, desc) in enumerate(tips, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i}. {title}\uff1a')
    set_run_font(run_num, size=Pt(11), bold=True, color=RGBColor(0xE6, 0x5C, 0x00))
    run_desc = p.add_run(desc)
    set_run_font(run_desc, size=Pt(10.5))

# --- Sources ---
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('\u53c2\u8003\u6765\u6e90\uff1a')
set_run_font(run, size=Pt(9), bold=True, color=RGBColor(0x99, 0x99, 0x99))
run2 = p.add_run('\u641c\u72d0\u201c\u903b\u8f91\u5b66\u201d\u786c\u6838\u751f\u5b58\u6307\u5357\u3001\u7f51\u6613\u957f\u671f\u50a8\u5b58\u98df\u54c1\u6307\u5357\u3001SharpCoderBlog\u6301\u4e45\u98df\u54c1\u6e05\u5355')
set_run_font(run2, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99))

# --- Save ---
output_path = r'f:\龙虾机器人\workbuddy生成的所有文件\可保存3-5年及以上的食物清单.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
